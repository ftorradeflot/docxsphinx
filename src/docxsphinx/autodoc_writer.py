import logging

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PAR_INDENT_PT = 20
SPACE_TITLE_CONTENT = 3

logger = logging.getLogger('docx')


class DescWriter():

    def __init__(self, translator, node, indent_level=0, par_indent=None):
        
        self.translator = translator
        self.parent_node = node
        self.node_type = node.get('objtype')
        self.desc_signature, self.desc_content = self.get_children_nodes_from_types(self.parent_node, \
            ['desc_signature', 'desc_content'])
        if par_indent:
            self.par_indent = par_indent
        else:
            self.par_indent = self.curr_indent
        self.or_indent = self.par_indent + indent_level*PAR_INDENT_PT

    def get_children_nodes_from_types(self, parent_node, types_list):
        
        children_nodes = []
        for node_type in types_list:
            chosen_node = None
            for child_node in parent_node.children:
                if child_node.tagname == node_type:
                    chosen_node = child_node
                    break
            children_nodes.append(chosen_node)
        return(children_nodes)
    
    def _add_text(self, text, new_paragraph=False, paragraph_format_kwargs=None, run_kwargs=None):
        
        if new_paragraph:
            p = self.translator.current_state.location.add_paragraph()
            self.translator.current_paragraph = p
        else:
            p = self.translator.current_paragraph
        r = p.add_run(text)
        
        if not run_kwargs is None:
            for k, v in run_kwargs.items():
                r.setattr(k, v)
        
        if not paragraph_format_kwargs is None:
            for k, v in paragraph_format_kwargs.items():
                p.paragraph_format.setattr(k, v)
    
    @property
    def curr_indent(self):
        curr_indent = self.translator.current_paragraph.paragraph_format.left_indent
        curr_indent = curr_indent and curr_indent.pt or 0
        return curr_indent
    
    def add_paragraph(self, text, bold=False, indent_level=0):
        # Create a new paragraph and tune it
        p = self.translator.current_state.location.add_paragraph()
        p.paragraph_format.left_indent = Pt(self.or_indent + indent_level*PAR_INDENT_PT)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        #Add the text
        r = p.add_run(text)
        r.bold = bold
        
        # Update the position and exit
        self.translator.current_paragraph = p
        return p
    
    def reset_indent(self):
        p = self.translator.current_state.location.add_paragraph()
        p.paragraph_format.left_indent = Pt(self.par_indent)
        self.translator.current_paragraph = p
    
    def write_signature(self):
        logger.info('Signature attributes: {}'.format(self.desc_signature.attributes))
        parameter_list, = self.get_children_nodes_from_types(self.desc_signature, ['desc_parameterlist'])
        args = parameter_list and ', '.join([node.astext() for node in parameter_list]) or ''
        if self.node_type in ['function', 'class', 'method']:
            args = '(' + args + ')'
        signature = self.node_type + ' ' +self.desc_signature['fullname'] + args
        
        # Add function call in bold
        p = self.add_paragraph(signature, bold=True)
        p.paragraph_format.space_after = Pt(SPACE_TITLE_CONTENT)

    def write_main(self):
        logger.info('Writing Function Main')
        logger.info('Function Desc signature = {}'.format(self.desc_signature))
        logger.info('Function Desc content = {}'.format(self.desc_content))
    
    
    def write_field_list(self, field_list_node):
        
        for field in field_list_node:
            
            p = self.add_paragraph(field.children[0].astext(), bold=True, indent_level=1)
            p.paragraph_format.space_after = Pt(SPACE_TITLE_CONTENT)
            
            p = self.add_paragraph(field.children[1].astext().replace('\n\n', '\n'), indent_level=2)
            p.paragraph_format.space_before = Pt(0)
    
    def write(self):
        self.write_main()
        logger.info('Indent previous to signature {}'.format(self.or_indent))
        self.write_signature()
        logger.info('Indent previous to content {}'.format(self.or_indent))
        self.write_content()
        logger.info('Indent after content {}'.format(self.or_indent))
        self.reset_indent()
    
    def write_content(self):
        
        for node in self.desc_content:

            if node.tagname == 'paragraph':
                self.add_paragraph(node.astext().replace('\n', ' '))
            
            elif node.tagname == 'field_list':
                self.write_field_list(node)
            
            elif node.tagname == 'desc':
                if node.get('objtype') == 'attribute':
                    w = AttributeWriter(self.translator, node,
                                        par_indent=self.or_indent,
                                        indent_level=1)
                else:
                    w = DescWriter(self.translator, node,
                                   par_indent=self.or_indent,
                                   indent_level=1)
                w.write()

class AttributeWriter(DescWriter):
        
    def write_content(self):
        paragraph, = self.get_children_nodes_from_types(self.desc_content, ['paragraph'])
        if paragraph:
            p = self.translator.current_paragraph
            r = p.add_run(' ' + paragraph.astext())
            r.bold = False        
        
    def write_signature(self):

        logger.info(self.desc_signature.attributes)
        
        desc_name, = self.get_children_nodes_from_types(self.desc_signature, ['desc_name'])
        signature = desc_name.astext()
        
        # Add function call in bold
        p = self.add_paragraph(signature, bold=True)
        logger.info(self.or_indent)
        p.paragraph_format.space_after = Pt(SPACE_TITLE_CONTENT)    

class ClassWriter(DescWriter):

    node_type = 'class'

    def write(self):
        self.write_main()
        self.write_signature()
        self.write_content()
        self.reset_indent()
                    
    def write_definition_list(self, definition_list_node):
        
        curr_loc = self.translator.current_state.location
        curr_indent = self.translator.current_paragraph.paragraph_format.left_indent
        curr_indent = curr_indent and curr_indent.pt or 0
        
        for definition_list_item in definition_list_node.children:
            #self.translator.add_text(definition_list_item.tagname + ', ' + definition_list_item.astext())
            #self.translator.add_text(str(definition_list_item.children))
            
            term, classifier, definition = self.get_children_nodes_from_types(definition_list_item,
                ['term', 'classifier', 'definition'])
            
            p = curr_loc.add_paragraph(term.astext() + ' : ' + classifier.astext())            
            p.paragraph_format.left_indent = Pt(curr_indent)
            p.paragraph_format.space_after = Pt(SPACE_TITLE_CONTENT)
            
            p = curr_loc.add_paragraph(definition.astext())
            p.paragraph_format.left_indent = Pt(curr_indent + PAR_INDENT_PT)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #self.translator.current_paragraph = p
            
